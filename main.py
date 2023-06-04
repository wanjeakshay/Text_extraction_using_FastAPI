from fastapi import FastAPI, File, UploadFile
from elasticsearch import Elasticsearch
import mysql.connector as mysql
from docx import Document
from typing import Union
import docx2txt
import PyPDF2
import shutil
import re
import os


app = FastAPI()

@app.post("/upload")
def upload(file: UploadFile = File(...)):

    # read pdf file.
    pdfReader = PyPDF2.PdfFileReader(file.file)
    totalpages = pdfReader.numPages

    d = {}
    for i in range(totalpages):
        pageObj = pdfReader.getPage(i)
        d[i] = pageObj.extractText()

    return {"filename": file.filename,"content": d}


@app.post("/doc_text_extraction")
def text_extract(file: UploadFile = File(...)):
    source_folder = r"C:\Users\Akshay\sample_500_2\\"
    destination_folder = r"C:\Users\Akshay\pythonProject1\\"

    lst = []
    d = {}

    # word document we want to open
    source = source_folder + file.filename
    destination = destination_folder + file.filename
    shutil.copy(source, destination)
    doc = Document(destination)
    for para in doc.paragraphs:
        lst.append(para.text)
        final = "".join(lst)
        d[0] = final

    return {'content': d}


@app.post("/pdf_text_extract")
async def text_extract(file: UploadFile = File(...)):

    location = r"C:\\Users\\Akshay Wanje\\pythonProject1\\"
    file_location = location + file.filename

    with open(file_location, "wb") as file_object:
        shutil.copyfileobj(file.file, file_object)

    head_tail = os.path.split(file.filename)

    lst =[]
    d ={}
    final = ''
    # see if file id pdf or not
    if head_tail[1][-4:] == '.pdf':
        # reading the pdf file
        pdfReader = PyPDF2.PdfFileReader(file_location)
        # number of pages
        totalpages = pdfReader.numPages

        for i in range(totalpages):
            pageObj = pdfReader.getPage(i)
            d[i] = pageObj.extractText()
            lst.append(d[i])
            d[i] = re.sub("\s+", " ", d[i])
        final = " ".join(lst)
        final = re.sub("\s+", " ", final)

    elif head_tail[1][-4:] == "docx":

        doc = Document(file_location)

        # extracting text line by line
        for para in doc.paragraphs:
            lst.append(para.text)
            final = " ".join(lst)
            final = re.sub("\s+", " ", final)
            d[0] = final

    else:
        pass

    db = mysql.connect(
        host="127.0.0.1",
        user="root", # add your mysql username
        passwd="password123", # add your mysql password
        database="user"
    )

    cursor = db.cursor()

    '''
    # below statement is used to create tha 'user' database
    cursor.execute("CREATE DATABASE user")
    
    # creating a table called 'candidate' in the 'user' database
    cursor.execute("CREATE TABLE candidate (candidate_id INT(11) NOT NULL AUTO_INCREMENT PRIMARY KEY, file_name VARCHAR(255), resume_content TEXT)")
    '''

    # defining the Query
    query = "INSERT INTO candidate (file_name, resume_content) VALUES (%s, %s)"
    # storing values in a variable
    values = (file.filename, final)

    # executing the query with values
    cursor.execute(query, values)

    # to make final output we have to run the 'commit()' method of the database object
    db.commit()

    return {"filename": file.filename, 'content': final, 'message':'Database is updated'}


@app.post("/update_elasticsearch")
async def update_elasticsearch(file: UploadFile = File(...)):
    location = r"C:\\Users\\Akshay Wanje\\PycharmProjects\\pythonProject1\\pythonProject1\\"
    file_location = location + file.filename

    with open(file_location, "wb") as file_object:
        shutil.copyfileobj(file.file, file_object)

    head_tail = os.path.split(file.filename)

    lst = []
    d = {}
    final = ''
    # see if file id pdf or not
    if head_tail[1][-4:] == '.pdf':
        # reading the pdf file
        pdfReader = PyPDF2.PdfFileReader(file_location)
        # number of pages
        totalpages = pdfReader.numPages

        for i in range(totalpages):
            pageObj = pdfReader.getPage(i)
            d[i] = pageObj.extractText()
            lst.append(d[i])
            d[i] = re.sub("\s+", " ", d[i])
        final = " ".join(lst)
        final = re.sub("\s+", " ", final)

    elif head_tail[1][-4:] == "docx":

        d = {}
        doc = Document(file_location)

        # extracting text line by line
        for para in doc.paragraphs:
            lst.append(para.text)
            final = " ".join(lst)
            final = re.sub("\s+", " ", final)
            d[0] = final

    else:
        pass

    # Connecting elasticsearch with python
    # Create the client instance with authentication.
    client = Elasticsearch("http://localhost:9200", http_auth=('elastic', 'fz9hj3x7KhFCAZ67D3-s')) # add your elasticsearch username password
    # Create an index.
    #client.indices.create(index="candidate_info")
    # update the index with data.
    client.index(index="candidate_info", id=head_tail[1], document=d)
    # to see a record.
    res = client.get(index='candidate_info', id=head_tail[1])

    return {"message":"Database is successfully updated",file.filename:res['_source']}
